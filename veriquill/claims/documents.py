"""Reading candidate-supplied documents into addressable lines.

Line numbers are the whole point: a claim has to be able to say "resume line
12", so blank lines are preserved and nothing is silently reflowed.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path


class UnsupportedDocument(RuntimeError):
    pass


@dataclass(frozen=True, slots=True)
class Document:
    name: str
    lines: tuple[str, ...]

    def line_number_of(self, text: str) -> int | None:
        """1-indexed line number of the first line equal to `text`."""
        needle = text.strip()
        for index, line in enumerate(self.lines, start=1):
            if line.strip() == needle:
                return index
        return None


def _read_text(path: Path) -> list[str]:
    raw = path.read_text(encoding="utf-8", errors="replace")
    return raw.splitlines()


def _read_pdf(path: Path) -> list[str]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    lines: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        lines.extend(text.splitlines())
    return lines


def _read_docx(path: Path) -> list[str]:
    import docx

    document = docx.Document(str(path))
    return [paragraph.text for paragraph in document.paragraphs]


_READERS = {
    ".txt": _read_text,
    ".md": _read_text,
    ".pdf": _read_pdf,
    ".docx": _read_docx,
}


def load_document(path: Path) -> Document:
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"document not found: {path}")

    reader = _READERS.get(path.suffix.lower())
    if reader is None:
        raise UnsupportedDocument(
            f"cannot read {path.suffix} documents; supported: "
            f"{', '.join(sorted(_READERS))}"
        )

    return Document(name=path.name, lines=tuple(reader(path)))
